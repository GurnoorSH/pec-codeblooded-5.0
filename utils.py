from bs4 import BeautifulSoup , NavigableString
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_horizontal_line(paragraph):
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert(0, pBdr)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.insert(0, bottom)
    paragraph.add_run('\n')

    
def parse_styles(css_style , tag_name):
    styles = {}
    
    if 'text-align:' in css_style:
        alignment = css_style.split('text-align:')[1].split(';')[0].strip()
        styles['alignment'] = alignment
    
    if 'font-weight:' in css_style:
        weight = css_style.split('font-weight:')[1].split(';')[0].strip()
        styles['bold'] = weight == 'bold'
    
    if 'font-style:' in css_style:
        font_style = css_style.split('font-style:')[1].split(';')[0].strip()
        styles['italic'] = font_style == 'italic'

    if 'font-size:' in css_style:
        size = int(css_style.split('font-size:')[1].split('px')[0])
        styles['size'] = size

    if 'color:' in css_style:
        color = css_style.split('color:')[1].split(';')[0].strip()
        if color.startswith('#'):
            color = RGBColor(int(color[1:3], 16), int(color[3:5], 16), int(color[5:7], 16))
            styles['color'] = color

    if 'size' not in styles:  # if size isn't specified in inline CSS
        if tag_name == 'h1':
            styles['size'] = 32
        elif tag_name == 'h2':
            styles['size'] = 24
        elif tag_name == 'h3':
            styles['size'] = 18.72
        elif tag_name == 'h4':
            styles['size'] = 16
        elif tag_name == 'h5':
            styles['size'] = 13.28
        elif tag_name == 'h6':
            styles['size'] = 10.72
    return styles

def html_to_docx(html_content):
    doc = Document()
    soup = BeautifulSoup(html_content, 'html.parser')

    for element in soup.find_all(True):
        if isinstance(element, NavigableString):
            continue
        
        if element.name == 'hr':
            para = doc.add_paragraph()
            add_horizontal_line(para)
            continue
        if element.name in ['style', 'script']:
            continue

        para = doc.add_paragraph()
        run = para.add_run(element.get_text())

        styles = parse_styles(element.get('style', ''), element.name)

        if 'alignment' in styles:
            if styles['alignment'] == 'center':
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif styles['alignment'] == 'right':
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif styles['alignment'] == 'left':
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if 'bold' in styles:
            run.bold = styles['bold']
        
        if 'italic' in styles:
            run.italic = styles['italic']

        if 'size' in styles:
            run.font.size = Pt(styles['size'])
        
        if 'color' in styles:
            run.font.color.rgb = styles['color']

    file_path = "output.docx"
    doc.save(file_path)
    return file_path
